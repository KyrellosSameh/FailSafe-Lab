# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ─── Global RTL + Arabic Font Setup ───
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
if rFonts is None:
    rPr = OxmlElement('w:rPr')
    style.element.append(rPr)
    rFonts = OxmlElement('w:rFonts')
    rPr.append(rFonts)
rFonts.set(qn('w:cs'), 'Arial')

# Make all heading styles RTL + Arabic
for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Arial'
    hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hRpr = hs.element.rPr
    if hRpr is None:
        hRpr = OxmlElement('w:rPr')
        hs.element.append(hRpr)
    hRf = hRpr.find(qn('w:rFonts'))
    if hRf is None:
        hRf = OxmlElement('w:rFonts')
        hRpr.append(hRf)
    hRf.set(qn('w:cs'), 'Arial')

def set_rtl(paragraph):
    """Set paragraph to RTL direction."""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def add_heading_rtl(text, level=1):
    h = doc.add_heading(text, level=level)
    set_rtl(h)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return h

def add_para(text, bold=False, size=12, color=None, align='right'):
    p = doc.add_paragraph()
    set_rtl(p)
    if align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    r.bold = bold
    rPr = r._element.get_or_add_rPr()
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:cs'), 'Arial')
    rPr.append(rF)
    rtl_elem = OxmlElement('w:rtl')
    rPr.append(rtl_elem)
    if color:
        r.font.color.rgb = color
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1E, 0x90, 0xFF)
    p.paragraph_format.left_indent = Cm(1)
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def add_table_rtl(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Arial'
        r.font.size = Pt(11)
        rPr = r._element.get_or_add_rPr()
        rF = OxmlElement('w:rFonts')
        rF.set(qn('w:cs'), 'Arial')
        rPr.append(rF)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.name = 'Arial'
            r.font.size = Pt(10)
            rPr = r._element.get_or_add_rPr()
            rF = OxmlElement('w:rFonts')
            rF.set(qn('w:cs'), 'Arial')
            rPr.append(rF)
    doc.add_paragraph()

# ════════════════════════════════════════════
#                DOCUMENT CONTENT
# ════════════════════════════════════════════

# ─── Title Page ───
doc.add_paragraph()
doc.add_paragraph()
add_para('FailSafe Lab', bold=True, size=36, color=RGBColor(0x00, 0x99, 0xCC), align='center')
add_para('الدليل الشامل للمشروع', bold=True, size=24, align='center')
add_para('شرح كامل لكل ملف في المشروع — للمبتدئين', size=14, color=RGBColor(0x66, 0x66, 0x66), align='center')
doc.add_paragraph()
add_para('منصة مختبر فيزياء افتراضية مع نظام امتحانات مؤمن', size=13, align='center')
doc.add_page_break()

# ─── Table of Contents ───
add_heading_rtl('فهرس المحتويات', 1)
toc_items = [
    'هيكل المشروع',
    'نقطة البداية (main.jsx)',
    'خريطة الطرق (App.jsx)',
    'الاتصال بقاعدة البيانات (supabaseClient.js)',
    'نظام الامتحانات (StudentRoutes.jsx)',
    'الإطار العام للمختبر (StudentLabLayout.jsx)',
    'نظام المراقبة بالكاميرا (CameraProctor.jsx)',
    'تجربة قانون أوم (OhmsLaw.jsx)',
    'تجربة قنطرة ويتستون (WheatstoneBridge.jsx)',
    'تجربة قانون هوك (HookesLaw.jsx)',
    'تجربة اللزوجة (Viscosity.jsx)',
    'صفحة دخول الطالب (StudentEntryPage.jsx)',
    'لوحة تحكم الأستاذ والأدمن',
    'قاعدة البيانات',
    'نظام الأمان',
    'التقنيات المستخدمة',
]
for i, item in enumerate(toc_items, 1):
    add_para(f'{i}. {item}', size=12)
doc.add_page_break()

# ═══════════════════════════════════════
# 1. هيكل المشروع
# ═══════════════════════════════════════
add_heading_rtl('1. هيكل المشروع', 1)
add_para('المشروع مبني بتقنية React وبيتكون من عدة مجلدات، كل مجلد له وظيفة محددة:')

add_heading_rtl('المجلدات الأساسية:', 2)
add_table_rtl(['المجلد', 'الوظيفة'], [
    ['src/', 'المجلد الرئيسي — كل كود المشروع'],
    ['src/pages/', 'الصفحات (الشاشات اللي بتظهر للمستخدم)'],
    ['src/features/', 'المميزات: التجارب + المراقبة'],
    ['src/routes/', 'ملفات التوجيه (مين يروح فين)'],
    ['src/layouts/', 'الهياكل المشتركة (sidebar + header)'],
    ['src/lib/', 'المكتبات المساعدة (اتصال قاعدة البيانات)'],
    ['src/styles/', 'ملفات التنسيق (CSS)'],
])

add_heading_rtl('الملفات داخل كل مجلد:', 2)
add_table_rtl(['الملف', 'الوظيفة', 'المستخدم'], [
    ['LandingPage.jsx', 'الصفحة الرئيسية', 'الكل'],
    ['StudentEntryPage.jsx', 'بوابة دخول الامتحان', 'الطالب'],
    ['OhmsLaw.jsx', 'تجربة قانون أوم', 'الطالب'],
    ['WheatstoneBridge.jsx', 'تجربة قنطرة ويتستون', 'الطالب'],
    ['HookesLaw.jsx', 'تجربة قانون هوك', 'الطالب'],
    ['Viscosity.jsx', 'تجربة اللزوجة', 'الطالب'],
    ['CameraProctor.jsx', 'نظام المراقبة بالذكاء الاصطناعي', 'الطالب (امتحان)'],
    ['InstructorDashboard.jsx', 'لوحة تحكم الأستاذ', 'الأستاذ'],
    ['CreateExamPage.jsx', 'إنشاء امتحان جديد', 'الأستاذ'],
    ['StudentResultsPage.jsx', 'عرض نتائج الطلاب', 'الأستاذ'],
    ['AdminDashboard.jsx', 'إدارة الطلاب', 'الأدمن'],
])
doc.add_page_break()

# ═══════════════════════════════════════
# 2. main.jsx
# ═══════════════════════════════════════
add_heading_rtl('2. نقطة البداية — main.jsx', 1)
add_para('هذا هو أول ملف يتم تنفيذه عند فتح الموقع. وظيفته ربط مكتبة React بصفحة HTML.')
add_code('import { StrictMode } from "react";\nimport { createRoot } from "react-dom/client";\nimport { BrowserRouter } from "react-router-dom";\nimport "./styles/core/index.css";\nimport App from "./App.jsx";')
add_para('')
add_heading_rtl('شرح كل سطر:', 2)
add_table_rtl(['الكود', 'الشرح'], [
    ['StrictMode', 'وضع صارم من React — يكشف الأخطاء أثناء التطوير فقط'],
    ['createRoot', 'الدالة التي تربط React بعنصر HTML الحقيقي (div#root)'],
    ['BrowserRouter', 'يفعّل نظام التنقل بين الصفحات عبر عناوين URL'],
    ['index.css', 'ملف التنسيقات الأساسي (الثيم، الألوان، الخطوط)'],
    ['App', 'المكوّن الرئيسي الذي يحتوي على كل التطبيق'],
])

add_code('createRoot(document.getElementById("root")).render(\n  <StrictMode>\n    <BrowserRouter>\n      <App />\n    </BrowserRouter>\n  </StrictMode>\n);')
add_para('التفسير: البحث عن عنصر HTML اسمه "root" ← رسم تطبيق React بالكامل بداخله.')
doc.add_page_break()

# ═══════════════════════════════════════
# 3. App.jsx
# ═══════════════════════════════════════
add_heading_rtl('3. خريطة الطرق — App.jsx', 1)
add_para('هذا الملف يعمل كموظف استقبال — يوجه كل زائر للمكان الصحيح بناءً على عنوان URL.')
add_code('<Routes>\n  <Route path="/" element={<LandingPage />} />\n  <Route path="/secure/ta-portal-4m8p1/*" element={<InstructorRoutes />} />\n  <Route path="/secure/ctrl-panel-9x7k2/*" element={<AdminRoutes />} />\n  <Route path="/lab/*" element={<StudentRoutes />} />\n  <Route path="*" element={<LandingPage />} />\n</Routes>')

add_table_rtl(['العنوان (URL)', 'يذهب إلى', 'المستخدم'], [
    ['/', 'الصفحة الرئيسية', 'الجميع'],
    ['/secure/ta-portal-4m8p1/', 'لوحة الأستاذ', 'الأستاذ فقط'],
    ['/secure/ctrl-panel-9x7k2/', 'لوحة الأدمن', 'الأدمن فقط'],
    ['/lab/', 'المختبر الافتراضي', 'الطالب'],
    ['أي عنوان آخر', 'الصفحة الرئيسية', 'إعادة توجيه'],
])
add_para('ملاحظة أمنية: عناوين الأستاذ والأدمن تحتوي على أكواد عشوائية (4m8p1, 9x7k2) كحماية — لا يمكن تخمينها بسهولة.', bold=True)
doc.add_page_break()

# ═══════════════════════════════════════
# 4. supabaseClient.js
# ═══════════════════════════════════════
add_heading_rtl('4. الاتصال بقاعدة البيانات — supabaseClient.js', 1)
add_para('Supabase هي خدمة سحابية توفر قاعدة بيانات + تخزين ملفات + دوال مؤمنة جاهزة للاستخدام.')
add_code('import { createClient } from "@supabase/supabase-js";\n\nconst supabaseUrl = import.meta.env.VITE_SUPABASE_URL;\nconst supabaseKey = import.meta.env.VITE_SUPABASE_ANON_KEY;\n\nexport const supabase = createClient(supabaseUrl, supabaseKey);')

add_table_rtl(['المتغير', 'الشرح'], [
    ['VITE_SUPABASE_URL', 'عنوان قاعدة البيانات (مخزن في ملف .env سري)'],
    ['VITE_SUPABASE_ANON_KEY', 'مفتاح الوصول العام'],
    ['createClient', 'إنشاء اتصال جاهز يُستخدم في كل أنحاء المشروع'],
])
add_para('في أي مكان نحتاج نقرأ أو نكتب بيانات، نستورد supabase ونستخدمها مباشرة.')
doc.add_page_break()

# ═══════════════════════════════════════
# 5. StudentRoutes.jsx
# ═══════════════════════════════════════
add_heading_rtl('5. نظام الامتحانات — StudentRoutes.jsx', 1)
add_para('هذا هو أهم ملف في المشروع. يتحكم في كل ما يخص الطالب: تسجيل الدخول، بدء الامتحان، العد التنازلي، الأمان، وتسليم الإجابة.', bold=True)

add_heading_rtl('المتغيرات الأساسية:', 2)
add_table_rtl(['المتغير', 'النوع', 'الوظيفة'], [
    ['activeTab', 'نص', 'التجربة النشطة حالياً (ohm, wheatstone, hooke, viscosity)'],
    ['examConfig', 'كائن أو null', 'إعدادات الامتحان الحالي — null يعني وضع تدريب'],
    ['timeLeft', 'رقم أو null', 'الوقت المتبقي بالميللي ثانية'],
    ['showWarning', 'صح/خطأ', 'هل نعرض تحذير آخر 5 دقائق؟'],
    ['isSubmitting', 'صح/خطأ', 'هل جاري التسليم؟ (لمنع التسليم المزدوج)'],
])

add_heading_rtl('نظام الأمان — منع الغش:', 2)
add_para('أثناء الامتحان، النظام يمنع الطالب من:')
add_table_rtl(['المفتاح/الإجراء', 'ماذا يمنع'], [
    ['كليك يمين', 'منع Inspect Element'],
    ['F12', 'منع فتح أدوات المطور'],
    ['Ctrl+Shift+I/J/C', 'منع اختصارات DevTools'],
    ['Ctrl+U', 'منع عرض كود المصدر'],
])

add_heading_rtl('تشفير بيانات الامتحان:', 2)
add_para('بيانات الامتحان تُحفظ مشفرة (Base64) في المتصفح، حيث أنه لو الطالب عمل تحديث (Refresh) للصفحة، الامتحان يكمل ولا يضيع.')
add_code('const safeBase64 = btoa(encodeURIComponent(JSON.stringify(examConfig)));\nlocalStorage.setItem("app_examConfig", safeBase64);')

add_heading_rtl('العد التنازلي (30 دقيقة):', 2)
add_para('كل ثانية بيحسب الوقت المتبقي:')
add_code('const elapsed = Date.now() - examConfig.startTime;\nconst remaining = Math.max(0, 30 * 60 * 1000 - elapsed);')
add_para('لو الوقت أقل من 5 دقائق ← تحذير أحمر متحرك')
add_para('لو الوقت = 0 ← تسليم تلقائي فوري')

add_heading_rtl('Heartbeat Ping (نبضة القلب):', 2)
add_para('كل 15 ثانية، النظام يبعث إشارة "أنا لسه موجود" لقاعدة البيانات. لو الإشارة توقفت، الأستاذ يعرف أن الطالب غادر.')

add_heading_rtl('تسليم الإجابة (handleExamSubmit):', 2)
add_para('دالة مركزية بتبعث الإجابة لقاعدة البيانات عبر دالة مؤمنة (RPC):')
add_table_rtl(['القيمة المرسلة', 'المعنى'], [
    ['قيمة عادية (مثل "24.5")', 'الطالب أجاب بشكل طبيعي'],
    ['"--"', 'انتهى الوقت — تسليم تلقائي'],
    ['"--quit--"', 'الطالب انسحب طواعية'],
    ['"--kicked--"', 'الكاميرا طردته لغياب الوجه'],
])
doc.add_page_break()

# ═══════════════════════════════════════
# 6. StudentLabLayout.jsx
# ═══════════════════════════════════════
add_heading_rtl('6. الإطار العام للمختبر — StudentLabLayout.jsx', 1)
add_para('هذا الملف يرسم الهيكل المشترك لكل التجارب: الشريط الجانبي (Sidebar) والهيدر ومنطقة التجربة.')

add_heading_rtl('التجارب الأربع:', 2)
add_table_rtl(['المعرّف (id)', 'الاسم', 'الأيقونة'], [
    ['ohm', "Ohm's Law (قانون أوم)", 'Activity ⚡'],
    ['wheatstone', 'Wheatstone (قنطرة ويتستون)', 'Atom ⚛️'],
    ['hooke', "Hooke's Law (قانون هوك)", 'Scale ⚖️'],
    ['viscosity', 'Viscosity (اللزوجة)', 'TestTube 🧪'],
])

add_heading_rtl('سلوك مختلف حسب الوضع:', 2)
add_table_rtl(['الوضع', 'الشريط الجانبي', 'الهيدر', 'الكاميرا'], [
    ['تصفح حر (Browse)', 'كل التجارب ظاهرة', 'اسم التجربة فقط', 'مغلقة'],
    ['امتحان (Exam)', 'تجربة واحدة فقط', 'اسم + ID + وقت متبقي', 'شغالة'],
])

add_para('في وضع الامتحان، زر الخروج يتحول إلى "انسحاب من الامتحان" باللون الأحمر مع رسالة تأكيد.')
doc.add_page_break()

# ═══════════════════════════════════════
# 7. CameraProctor.jsx
# ═══════════════════════════════════════
add_heading_rtl('7. نظام المراقبة بالكاميرا — CameraProctor.jsx', 1)
add_para('من أهم وأخطر الملفات في المشروع — يتحكم في مراقبة الطالب بالكاميرا ويمكنه طرده من الامتحان!', bold=True, color=RGBColor(0xCC, 0x00, 0x00))

add_heading_rtl('آلية العمل:', 2)
add_para('1. عند بدء الامتحان ← يفتح الكاميرا')
add_para('2. كل 3 ثوانٍ ← يفحص هل الوجه موجود؟')
add_para('3. لو الوجه غائب 4 مرات متتالية (12 ثانية) ← تحذير!')
add_para('4. الطالب عنده 30 ثانية يرجع أمام الكاميرا')
add_para('5. لو ما رجعش ← طرد تلقائي + إبلاغ الأستاذ + صورة')

add_heading_rtl('كشف الوجه (face-api.js):', 2)
add_para('يستخدم مكتبة ذكاء اصطناعي (شبكة عصبية مدربة) لكشف الوجوه في الفيديو:')
add_code('const detection = await window.faceapi.detectSingleFace(\n  videoRef.current,\n  new window.faceapi.TinyFaceDetectorOptions({\n    inputSize: 320,        // حجم أكبر = دقة أعلى\n    scoreThreshold: 0.25,  // حد أدنى منخفض = تسامح أكثر\n  })\n);')

add_heading_rtl('التقاط الصور:', 2)
add_table_rtl(['النوع', 'التكرار', 'المكان في التخزين'], [
    ['صورة دورية', 'كل دقيقة', 'exam_snapshots/{code}/{student_id}/'],
    ['صورة تنبيه', 'عند الطرد', 'exam_snapshots/alerts/{code}/{student_id}/'],
])

add_heading_rtl('كشف الخروج من المتصفح:', 2)
add_table_rtl(['الحدث', 'النوع', 'الإبلاغ'], [
    ['الطالب بدّل لتبويب آخر', 'tab_switch', 'إشعار فوري للأستاذ'],
    ['الطالب أغلق المتصفح', 'browser_closed', 'إشعار عبر sendBeacon'],
    ['الطالب غاب عن الكاميرا', 'face_absent', 'إشعار + صورة'],
])

add_para('تقنية sendBeacon: واجهة برمجية خاصة تضمن إرسال البيانات حتى لو الصفحة تُغلق — على عكس fetch العادي الذي يتوقف.', bold=True)
doc.add_page_break()

# ═══════════════════════════════════════
# 8-11. التجارب الأربعة
# ═══════════════════════════════════════
add_heading_rtl('8. تجربة قانون أوم — OhmsLaw.jsx', 1)
add_para('القانون: R = V / I', bold=True, size=16)
add_para('الهدف: قياس الجهد (V) والتيار (I) وحساب المقاومة المجهولة (R).')
add_heading_rtl('المكونات:', 2)
add_table_rtl(['المكوّن', 'الشرح'], [
    ['الفولتميتر', 'شاشة ديجيتال تعرض الجهد — مصممة كجهاز حقيقي بأرقام حمراء'],
    ['الأميتير', 'شاشة ديجيتال تعرض التيار'],
    ['شريط الجهد (Slider)', 'يتحرك لتغيير جهد مصدر الطاقة (0-24V)'],
    ['المقاومة المجهولة', 'مرسومة بألوان حسب قيمتها'],
    ['صندوق الحساب', 'الطالب يكتب إجابته هنا'],
])
add_para('المنطق: المقاومة = رقم عشوائي مخفي. الطالب يغير الجهد ← يقرأ التيار ← يحسب R = V / I')
doc.add_page_break()

add_heading_rtl('9. تجربة قنطرة ويتستون — WheatstoneBridge.jsx', 1)
add_para('القانون: Rx = R × (L / (100 - L))', bold=True, size=16)
add_para('الهدف: إيجاد المقاومة المجهولة Rx عن طريق موازنة الجسر.')
add_heading_rtl('المكونات:', 2)
add_table_rtl(['المكوّن', 'الشرح'], [
    ['الجلفانوميتر', 'مؤشر تناظري. الهدف جعله على الصفر'],
    ['السلك المتري (Slider)', 'تحريك "الجوكي" على السلك (0-100cm)'],
    ['المقاومة المعروفة R', 'تتغير (100Ω, 200Ω, ...)'],
    ['المقاومة المجهولة Rx', 'ما يحسبه الطالب'],
])
add_para('مبدأ التوازن: عندما المؤشر = 0 عند موقع L ← Rx = R × (L / (100 - L))')
doc.add_page_break()

add_heading_rtl('10. تجربة قانون هوك — HookesLaw.jsx', 1)
add_para('القانون: k = F / x', bold=True, size=16)
add_para('الهدف: حساب ثابت مرونة الزنبرك المجهول k.')
add_heading_rtl('المكونات:', 2)
add_table_rtl(['المكوّن', 'الشرح'], [
    ['الزنبرك', 'مرسوم بـ SVG Bezier Curves — يتمدد واقعياً'],
    ['المسطرة', '51 علامة ذهبية (0-50cm)'],
    ['صينية الأوزان', '5 أوزان ملونة يمكن سحبها وإفلاتها'],
    ['القراءة الحية', 'F (قوة) و x (إزاحة) بالوقت الحقيقي'],
    ['محرك الأنيميشن', 'فيزياء حقيقية: Spring-Mass-Damper'],
])

add_heading_rtl('محرك الأنيميشن — التخميد الديناميكي:', 2)
add_code('const damping = 0.25 * 2 * Math.sqrt(m * k);')
add_para('هذه المعادلة تحسب التخميد كنسبة 25% من التخميد الحرج. النتيجة: الزنبرك يهتز بشكل واقعي (3-4 اهتزازات) ثم يستقر، بغض النظر عن حجم الوزن.')
doc.add_page_break()

add_heading_rtl('11. تجربة اللزوجة — Viscosity.jsx', 1)
add_para('القانون: η = (2r²(ρs - ρf)g) / (9v)', bold=True, size=16)
add_para('الهدف: قياس لزوجة سائل مجهول عن طريق إسقاط كرات فيه.')
add_heading_rtl('المكونات:', 2)
add_table_rtl(['المكوّن', 'الشرح'], [
    ['الأنبوبة', 'أنبوبة طويلة فيها سائل مرسومة بـ CSS'],
    ['5 كرات', 'كل واحدة بقطر وكثافة مختلفة'],
    ['الميكروميتر', 'لقياس قطر الكرة يدوياً'],
    ['الساعة (Stopwatch)', 'لقياس وقت سقوط الكرة بين خطين'],
])

add_heading_rtl('خطوات التجربة:', 2)
add_para('1. اختيار كرة من الكرات الخمسة')
add_para('2. قياس قطرها بالميكروميتر (المقياس الرئيسي + المقياس الدائري)')
add_para('3. إسقاطها في الأنبوبة')
add_para('4. الضغط على Start عندما تعبر الخط الأول')
add_para('5. الضغط على Stop عندما تعبر الخط الثاني')
add_para('6. حساب السرعة = المسافة ÷ الزمن')
add_para('7. حساب اللزوجة بالمعادلة')
doc.add_page_break()

# ═══════════════════════════════════════
# 12. StudentEntryPage
# ═══════════════════════════════════════
add_heading_rtl('12. صفحة دخول الطالب — StudentEntryPage.jsx', 1)
add_para('البوابة الرئيسية التي يدخل منها الطالب للامتحان.')

add_heading_rtl('تدفق الدخول:', 2)
add_para('1. فحص المتصفح — هل هو Safe Exam Browser (SEB)؟')
add_para('   - لو لا ← تحذير أحمر مع رابط تحميل SEB')
add_para('   - لو نعم ← متابعة')
add_para('2. إدخال الرقم الأكاديمي + كلمة المرور')
add_para('3. مقارنة كلمة المرور بالنسخة المشفرة (bcrypt)')
add_para('4. التحقق من حالة الطالب:')
add_para('   - لو عنده امتحان معلق ← الدخول مباشرة')
add_para('   - لو أكمل الامتحان ← رسالة "لقد أكملت اختبارك"')
add_para('   - لو مقفول (خرج من SEB) ← "تم قفل الامتحان نهائياً"')
add_para('   - لو مفيش امتحان ← "سحب اختبار عشوائي"')

add_heading_rtl('تشفير كلمة المرور:', 2)
add_code('import bcrypt from "bcryptjs";\nconst isMatch = await bcrypt.compare(studentPassword, hashedPassword);')
add_para('كلمات المرور مخزنة مشفرة في قاعدة البيانات باستخدام خوارزمية bcrypt. عند المقارنة لا يتم فك التشفير بل تتم مقارنة مشفرة.')
doc.add_page_break()

# ═══════════════════════════════════════
# 13. لوحات التحكم
# ═══════════════════════════════════════
add_heading_rtl('13. لوحات تحكم الأستاذ والأدمن', 1)

add_heading_rtl('لوحة الأستاذ (InstructorDashboard.jsx):', 2)
add_table_rtl(['القسم', 'الوظيفة'], [
    ['بنك الامتحانات', 'عرض الامتحانات المنشأة + إدارتها'],
    ['تنبيهات المراقبة', 'عرض تنبيهات: تبديل تبويب، غياب الوجه، إغلاق المتصفح'],
    ['نتائج الطلاب', 'جدول بكل النتائج مع إمكانية التصدير PDF'],
    ['التقييم', 'إعطاء درجات للطلاب وحفظها في قاعدة البيانات'],
])

add_heading_rtl('لوحة الأدمن (AdminDashboard.jsx):', 2)
add_table_rtl(['الوظيفة', 'الشرح'], [
    ['إضافة طالب', 'اسم + ID + بريد + كلمة سر مشفرة بـ bcrypt'],
    ['حذف طالب', 'إزالة من قاعدة البيانات'],
    ['البحث والفلترة', 'البحث بالاسم أو الرقم الأكاديمي'],
    ['إعادة تعيين كلمة السر', 'توليد كلمة سر جديدة مشفرة'],
])
doc.add_page_break()

# ═══════════════════════════════════════
# 14. قاعدة البيانات
# ═══════════════════════════════════════
add_heading_rtl('14. قاعدة البيانات (Supabase)', 1)

add_heading_rtl('الجداول:', 2)
add_table_rtl(['الجدول', 'الوظيفة', 'أهم الأعمدة'], [
    ['students', 'بيانات الطلاب', 'student_id, name, email, password_hash'],
    ['exams', 'الامتحانات', 'session_code, experiment, parameters, student_id, opened_at'],
    ['results', 'نتائج الطلاب', 'student_id, exam_code, student_result, actual_result, unit'],
    ['proctor_alerts', 'تنبيهات المراقبة', 'student_id, exam_code, alert_type, snapshot_path'],
])

add_heading_rtl('التخزين (Storage):', 2)
add_table_rtl(['المجلد', 'المحتوى'], [
    ['exam_snapshots/', 'صور الطلاب الدورية أثناء الامتحان'],
    ['exam_snapshots/alerts/', 'صور التنبيهات (غياب الوجه)'],
])
doc.add_page_break()

# ═══════════════════════════════════════
# 15. نظام الأمان
# ═══════════════════════════════════════
add_heading_rtl('15. نظام الأمان — 10 طبقات حماية', 1)
add_para('المشروع يعتمد على 10 طبقات حماية متكاملة لمنع الغش أثناء الامتحان:', bold=True)

add_table_rtl(['الطبقة', 'ضد ماذا؟', 'كيف تعمل؟'], [
    ['1. Safe Exam Browser', 'فتح تطبيقات/تبويبات أخرى', 'متصفح خاص يقفل الجهاز أثناء الامتحان'],
    ['2. كاميرا AI', 'غياب الطالب', 'كشف الوجه كل 3 ثوانٍ + طرد بعد 30 ثانية غياب'],
    ['3. منع DevTools', 'التلاعب بالكود', 'حظر F12, Ctrl+Shift+I, كليك يمين'],
    ['4. تشفير Base64', 'التلاعب ببيانات الامتحان', 'بيانات الامتحان مشفرة في localStorage'],
    ['5. Heartbeat كل 15 ثانية', 'انقطاع مجهول', 'إشارة "أنا موجود" لقاعدة البيانات'],
    ['6. كشف تبديل التبويب', 'فتح تبويب آخر للغش', 'إبلاغ فوري عبر visibilitychange'],
    ['7. كشف إغلاق المتصفح', 'الخروج ثم العودة', 'إبلاغ عبر sendBeacon + قفل الامتحان'],
    ['8. RPC Functions', 'التلاعب بقاعدة البيانات', 'التسليم عبر دوال مؤمنة في الخادم'],
    ['9. قفل بعد الفتح', 'الخروج والعودة', 'بمجرد فتح الامتحان + الخروج = قفل نهائي'],
    ['10. وقت محدود (30 دقيقة)', 'المماطلة', 'تسليم تلقائي عند انتهاء الوقت'],
])
doc.add_page_break()

# ═══════════════════════════════════════
# 16. التقنيات
# ═══════════════════════════════════════
add_heading_rtl('16. التقنيات المستخدمة', 1)
add_table_rtl(['التقنية', 'الاستخدام'], [
    ['React 18', 'بناء واجهة المستخدم (المكونات التفاعلية)'],
    ['React Router v6', 'التنقل بين الصفحات بدون إعادة تحميل'],
    ['Vite', 'أداة التطوير والبناء (سريعة جداً)'],
    ['Supabase', 'قاعدة بيانات + تخزين ملفات + RPC'],
    ['bcryptjs', 'تشفير كلمات المرور'],
    ['face-api.js', 'كشف الوجوه بالذكاء الاصطناعي (شبكة عصبية)'],
    ['EmailJS', 'إرسال إيميلات (استعادة كلمة السر)'],
    ['lucide-react', 'أيقونات SVG جاهزة'],
    ['CSS Variables', 'نظام الثيم (تغيير الألوان مركزياً)'],
    ['SVG Paths + Bezier Curves', 'رسم الزنبرك والجلفانوميتر'],
    ['Canvas API', 'التقاط صور من كاميرا الويب'],
    ['localStorage', 'حفظ بيانات الجلسة محلياً'],
    ['navigator.sendBeacon', 'إرسال بيانات عند إغلاق المتصفح'],
])

doc.add_paragraph()
add_para('— نهاية الدليل —', bold=True, size=14, align='center', color=RGBColor(0x00, 0x99, 0xCC))

# ─── Save ───
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'FailSafe_Lab_Documentation.docx')
doc.save(output_path)
print(f"Document saved to: {output_path}")
